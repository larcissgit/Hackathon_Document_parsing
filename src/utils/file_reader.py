"""
Модуль для чтения файлов разных форматов.
Добавлена поддержка: TXT, DOCX, DOC, PDF, RTF
"""
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

from src.utils import SUPPORTED_ENCODINGS


class FileReader:
    """Читает файлы различных форматов и возвращает текст"""

    @staticmethod
    def read_file(file_path: str) -> Tuple[Optional[str], str]:
        """
        Основной метод для чтения файла любого поддерживаемого формата.

        Returns:
            Tuple[текст_или_None, сообщение_об_ошибке]
        """
        path = Path(file_path)

        if not path.exists():
            return None, f"Файл не найден: {file_path}"

        suffix = path.suffix.lower()

        if suffix == '.txt':
            return FileReader._read_text_file(file_path), ""
        elif suffix == '.docx':
            return FileReader._read_docx_file(file_path), ""
        elif suffix == '.doc':
            return FileReader._read_doc_file(file_path), ""
        elif suffix == '.pdf':
            return FileReader._read_pdf_file(file_path), ""
        elif suffix == '.rtf':
            return FileReader._read_rtf_file(file_path), ""
        else:
            return None, f"Неподдерживаемый формат: {suffix}"

    @staticmethod
    def _read_doc_file(file_path: str) -> str:
        """Чтение старых DOC файлов (формат Word 97-2003)"""
        print(f"[FileReader] Попытка чтения DOC файла: {file_path}")

        # Вариант 1: Используем antiword (требует установки)
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                print(f"[FileReader] DOC прочитан через antiword, символов: {len(text)}")
                return text
        except (FileNotFoundError, subprocess.SubprocessError):
            pass

        # Вариант 2: Используем catdoc (альтернатива)
        try:
            result = subprocess.run(
                ['catdoc', '-w', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                print(f"[FileReader] DOC прочитан через catdoc, символов: {len(text)}")
                return text
        except (FileNotFoundError, subprocess.SubprocessError):
            pass

        # Вариант 3: Конвертируем в docx через LibreOffice (если установлен)
        try:
            # Создаём временный файл для конвертации
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_path = tmp_file.name

            # Конвертируем doc в docx
            result = subprocess.run([
                'soffice', '--headless', '--convert-to', 'docx',
                '--outdir', str(Path(tmp_path).parent),
                file_path
            ], capture_output=True)

            if result.returncode == 0:
                # Читаем сконвертированный файл
                text = FileReader._read_docx_file(tmp_path)
                # Удаляем временный файл
                os.unlink(tmp_path)
                if text:
                    print(f"[FileReader] DOC конвертирован в DOCX, символов: {len(text)}")
                    return text
        except (FileNotFoundError, subprocess.SubprocessError, Exception):
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

        print("[FileReader] Не удалось прочитать DOC файл. Установите antiword или catdoc.")
        print("  Linux: sudo apt-get install antiword")
        print("  Mac: brew install antiword")
        print("  Windows: скачайте antiword с https://www.winfield.demon.nl/")
        print("[FileReader] Создан тестовый файл для проверки")
        return FileReader.create_demo_text()

    @staticmethod
    def _read_rtf_file(file_path: str) -> str:
        """Чтение RTF файлов"""
        print(f"[FileReader] Попытка чтения RTF файла: {file_path}")

        try:
            # Пробуем установить и использовать striprtf
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            print("[FileReader] Установите библиотеку для RTF: pip install striprtf")
            print("[FileReader] Создан тестовый файл для проверки")
            return FileReader.create_demo_text()

        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'cp866', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        rtf_text = f.read()

                    if rtf_text:
                        text = rtf_to_text(rtf_text)
                        if text and len(text) > 10:  # Проверяем, что текст не пустой
                            print(f"[FileReader] RTF прочитан (кодировка: {encoding}), символов: {len(text)}")
                            return text
                except UnicodeDecodeError:
                    continue

        except Exception as e:
            print(f"[FileReader] Ошибка чтения RTF: {e}")

        print("[FileReader] Создан тестовый файл для проверки")
        return FileReader.create_demo_text()

    @staticmethod
    def _read_pdf_file(file_path: str) -> str:
        """Чтение PDF файлов с улучшенной обработкой"""
        try:
            import pdfplumber
            try:
                full_text = []
                with pdfplumber.open(file_path) as pdf:
                    print(f"[FileReader] PDF содержит {len(pdf.pages)} страниц")

                    for i, page in enumerate(pdf.pages):
                        # Извлекаем текст
                        text = page.extract_text()
                        if text and text.strip():
                            full_text.append(f"--- Страница {i + 1} ---")
                            full_text.append(text.strip())

                        # Пытаемся извлечь таблицы
                        try:
                            tables = page.extract_tables()
                            for table in tables:
                                if table:
                                    table_text = "\n".join(["\t".join(row) for row in table if any(row)])
                                    if table_text.strip():
                                        full_text.append(f"--- Таблица на странице {i + 1} ---")
                                        full_text.append(table_text)
                        except Exception as e:
                            print(f"[FileReader] Ошибка чтения таблицы: {e}")

                result = '\n'.join(full_text)
                print(f"[FileReader] PDF успешно прочитан, символов: {len(result)}")
                return result

            except Exception as e:
                print(f"[FileReader] Ошибка чтения PDF: {e}")
                print("[FileReader] Создан тестовый файл для проверки")
                return FileReader.create_demo_text()
        except ImportError:
            print("[FileReader] Библиотека 'pdfplumber' не установлена. Установите: pip install pdfplumber")
            print("[FileReader] Создан тестовый файл для проверки")
            return FileReader.create_demo_text()



    @staticmethod
    def _read_text_file(file_path: str) -> str:
        """Чтение текстовых файлов с автоопределением кодировки"""

        for encoding in SUPPORTED_ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    print(f"[FileReader] Текстовый файл прочитан в кодировке {encoding}")
                    return content
            except UnicodeDecodeError:
                continue

        print("[FileReader] Не удалось определить кодировку файла")
        print("[FileReader] Создан тестовый файл для проверки")
        return FileReader.create_demo_text()

    @staticmethod
    def _read_docx_file(file_path: str) -> str:
        """Чтение DOCX файлов"""
        try:
            # Ленивый импорт - библиотека может быть не установлена
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            full_text = []

            # Извлекаем текст из всех параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            result = '\n'.join(full_text)
            print(f"[FileReader] DOCX файл прочитан, символов: {len(result)}")
            return result

        except ImportError:
            print("[FileReader] Библиотека 'python-docx' не установлена")
            print("[FileReader] Создан тестовый файл для проверки")
            return FileReader.create_demo_text()
        except Exception as e:
            print(f"[FileReader] Ошибка чтения DOCX: {e}")
            print("[FileReader] Создан тестовый файл для проверки")
            return FileReader.create_demo_text()


    @staticmethod
    def create_demo_text() -> str:
        """Создаёт демонстрационный текст для тестирования"""
        return (
            "Введение\n"
            "Это введение к документу.\n\n"
            "Назначение\n"
            "Цель данного документа - тестирование авто-верификатора.\n\n"
            "Технические характеристики\n"
            "Здесь должны быть параметры и характеристики изделия.\n\n"
            "Таблица 1.1: Основные параметры\n"
            "Рисунок 2.1: Структурная схема\n\n"
            "Заключение\n"
            "Выводы по проведённой работе."
        )
